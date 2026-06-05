from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "generated"
OUT_FILE = OUT_DIR / "iptk_xulosasi.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = "DDE5DE", size: str = "8") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=90, start=150, bottom=90, end=150) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for key, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = margins.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_pct: int = 5000) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_pct))
    tbl_w.set(qn("w:type"), "pct")


def style_run(run, size=10.5, bold=False, color="1F2937") -> None:
    run.font.name = "Roboto"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Roboto")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_label_value_table(doc: Document, rows, widths=(4.2, 11.2)):
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_table_width(table)
    for label, value in rows:
        row = table.add_row()
        row.cells[0].width = Cm(widths[0])
        row.cells[1].width = Cm(widths[1])
        for cell in row.cells:
            set_cell_border(cell)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
        set_cell_shading(row.cells[0], "F3F7F3")
        label_run = row.cells[0].paragraphs[0].add_run(label)
        style_run(label_run, size=9.5, bold=True, color="6B7280")
        value_run = row.cells[1].paragraphs[0].add_run(value)
        style_run(value_run, size=10, bold=False, color="111827")
    return table


def add_section_title(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(9)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    style_run(run, size=11.5, bold=True, color="0F172A")


def add_note_box(doc: Document, title: str, body: str, fill="F6FAF6", border="B7D7BA") -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, border)
    set_cell_margins(cell, top=140, start=200, bottom=140, end=200)
    title_run = cell.paragraphs[0].add_run(title)
    style_run(title_run, size=10.5, bold=True, color="166534")
    paragraph = cell.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    body_run = paragraph.add_run(body)
    style_run(body_run, size=10, color="1F2937")


def add_signature_table(doc: Document) -> None:
    table = doc.add_table(rows=3, cols=3)
    set_table_width(table)
    headers = ("Lavozim", "F.I.O.", "Imzo")
    rows = (
        ("Komissiya raisi", "ABDULLAYEV KOMILJON ZAFAROVICH", "____________"),
        ("Kotib", "KARIMOVA DILNOZA BAHROM QIZI", "____________"),
    )
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_shading(cell, "EEF5EE")
        set_cell_border(cell)
        set_cell_margins(cell)
        run = cell.paragraphs[0].add_run(header)
        style_run(run, size=9.5, bold=True, color="64748B")
    for row_idx, data in enumerate(rows, start=1):
        for col_idx, value in enumerate(data):
            cell = table.cell(row_idx, col_idx)
            set_cell_border(cell)
            set_cell_margins(cell)
            run = cell.paragraphs[0].add_run(value)
            style_run(run, size=9.5, bold=col_idx == 1, color="111827")


def build_docx() -> None:
    OUT_DIR.mkdir(exist_ok=True)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.15)
    section.bottom_margin = Cm(1.15)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    styles = doc.styles
    styles["Normal"].font.name = "Roboto"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Roboto")
    styles["Normal"].font.size = Pt(10)

    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.space_after = Pt(4)
    run = heading.add_run("IJTIMOIY-PSIXIATRIK-TIBBIY KOMISSIYA XULOSASI")
    style_run(run, size=14, bold=True, color="0F172A")

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(9)
    run = sub.add_run("IPTK-XUL-2026-001-son")
    style_run(run, size=10.5, bold=True, color="4B5563")

    meta_table = doc.add_table(rows=1, cols=3)
    set_table_width(meta_table)
    meta = (
        ("Sana", "13.05.2026"),
        ("Hudud", "Buxoro viloyati"),
        ("Natija", "Ijobiy"),
    )
    for i, (label, value) in enumerate(meta):
        cell = meta_table.cell(0, i)
        set_cell_shading(cell, "F8FAF8")
        set_cell_border(cell)
        set_cell_margins(cell, top=80, start=140, bottom=80, end=140)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        label_run = p.add_run(f"{label}\n")
        style_run(label_run, size=8.5, bold=True, color="64748B")
        value_run = p.add_run(value)
        style_run(value_run, size=10.5, bold=True, color="111827")

    add_section_title(doc, "1. Xizmat oluvchi ma'lumotlari")
    add_label_value_table(
        doc,
        [
            ("F.I.O.", "SAIDOVA NILUFAR AKMAL QIZI"),
            ("Tug'ilgan sanasi", "06.06.2008"),
            ("JSHSHIR", "10**********85"),
            ("Nogironlik guruhi", "II guruh"),
            ("Tashxis", "Og'ir darajadagi aqliy zaiflik (F72)"),
            ("Manzil", "Buxoro viloyati, Buxoro shahri"),
        ],
    )

    add_section_title(doc, "2. Ariza va xizmat ma'lumotlari")
    add_label_value_table(
        doc,
        [
            ("Ariza raqami", "ARZ-000006"),
            ("Ariza sanasi", "06.06.2026"),
            ("Tanlangan xizmat", "G'amxo'rlik markaziga joylashtirish"),
            ("Tavsiya muddati", "12 oy"),
            ("Ko'rib chiqish asosi", "IPTK yig'ilishi bayonnomasi va biriktirilgan hujjatlar"),
        ],
    )

    add_section_title(doc, "3. Komissiya xulosasi")
    add_note_box(
        doc,
        "Xulosa",
        "Komissiya tomonidan ko'rib chiqilgan ma'lumotlar va hujjatlar asosida xizmat oluvchini "
        "tanlangan ijtimoiy xizmatga yo'naltirish maqsadga muvofiq deb topildi.",
    )

    add_section_title(doc, "4. Qaror")
    decision_table = doc.add_table(rows=1, cols=2)
    set_table_width(decision_table)
    for cell in decision_table.row_cells(0):
        set_cell_border(cell)
        set_cell_margins(cell, top=120, start=200, bottom=120, end=200)
    set_cell_shading(decision_table.cell(0, 0), "ECFDF5")
    set_cell_shading(decision_table.cell(0, 1), "F8FAF8")
    run = decision_table.cell(0, 0).paragraphs[0].add_run("Ijobiy xulosa")
    style_run(run, size=11, bold=True, color="047857")
    run = decision_table.cell(0, 1).paragraphs[0].add_run(
        "G'amxo'rlik markaziga joylashtirish tavsiya etiladi."
    )
    style_run(run, size=10.5, color="111827")

    add_section_title(doc, "5. Tasdiqlash")
    add_signature_table(doc)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("Ushbu hujjat axborot tizimi orqali shakllantirildi.")
    style_run(footer_run, size=8.5, color="6B7280")

    doc.save(OUT_FILE)


if __name__ == "__main__":
    build_docx()
    print(OUT_FILE)
